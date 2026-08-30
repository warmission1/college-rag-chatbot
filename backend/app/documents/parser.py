import io
import os
import re
from typing import List, Optional
from pydantic import BaseModel
from pypdf import PdfReader
import docx
from backend.app.core.errors import DocumentInvalidError


class ExtractedPage(BaseModel):
    page_number: int
    text: str
    headings: List[str] = []


class ExtractedDocument(BaseModel):
    title: Optional[str] = None
    page_count: int
    pages: List[ExtractedPage]
    warnings: List[str] = []
    ocr_used: bool = False
    parser_version: str = "1.0.0"


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf_bytes(file_bytes: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages: List[ExtractedPage] = []
        warnings = []
        
        if len(reader.pages) == 0:
            raise DocumentInvalidError("PDF contains zero pages")
            
        for idx, page in enumerate(reader.pages, start=1):
            raw_text = page.extract_text() or ""
            cleaned = clean_text(raw_text)
            
            headings = []
            for line in cleaned.split("\n"):
                s = line.strip()
                if s.startswith("#") or (len(s) > 3 and len(s) < 60 and s.isupper()):
                    headings.append(s.lstrip("#").strip())
                    
            if not cleaned:
                warnings.append(f"Page {idx} text was empty or scanned image without embedded text.")
                
            pages.append(ExtractedPage(page_number=idx, text=cleaned, headings=headings))
            
        return ExtractedDocument(
            page_count=len(pages),
            pages=pages,
            warnings=warnings,
            ocr_used=False,
        )
    except Exception as exc:
        if isinstance(exc, DocumentInvalidError):
            raise exc
        raise DocumentInvalidError(f"Failed to parse PDF document: {str(exc)}")


def parse_docx_bytes(file_bytes: bytes) -> ExtractedDocument:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        page_text_blocks = []
        headings = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                headings.append(text)
                page_text_blocks.append(f"\n## {text}\n")
            else:
                page_text_blocks.append(text)
                
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                page_text_blocks.append(" | ".join(row_cells))
                
        full_text = clean_text("\n".join(page_text_blocks))
        if not full_text:
            raise DocumentInvalidError("DOCX file is empty")
            
        pages = [ExtractedPage(page_number=1, text=full_text, headings=headings)]
        return ExtractedDocument(
            page_count=1,
            pages=pages,
            warnings=[],
            ocr_used=False,
        )
    except Exception as exc:
        if isinstance(exc, DocumentInvalidError):
            raise exc
        raise DocumentInvalidError(f"Failed to parse DOCX document: {str(exc)}")


def parse_text_bytes(file_bytes: bytes) -> ExtractedDocument:
    try:
        content = file_bytes.decode("utf-8", errors="ignore")
        cleaned = clean_text(content)
        if not cleaned:
            raise DocumentInvalidError("Text document is empty")
            
        headings = []
        for line in cleaned.split("\n"):
            s = line.strip()
            if s.startswith("#") or (s.isupper() and 3 < len(s) < 60):
                headings.append(s.lstrip("#").strip())
                
        paragraphs = cleaned.split("\n\n")
        pages = []
        current_page_text = []
        current_word_count = 0
        page_num = 1
        
        for para in paragraphs:
            current_page_text.append(para)
            current_word_count += len(para.split())
            if current_word_count >= 600:
                pages.append(ExtractedPage(
                    page_number=page_num,
                    text="\n\n".join(current_page_text),
                    headings=[h for h in headings if h in "\n\n".join(current_page_text)],
                ))
                current_page_text = []
                current_word_count = 0
                page_num += 1
                
        if current_page_text:
            pages.append(ExtractedPage(
                page_number=page_num,
                text="\n\n".join(current_page_text),
                headings=[h for h in headings if h in "\n\n".join(current_page_text)],
            ))
            
        return ExtractedDocument(
            page_count=len(pages),
            pages=pages,
            warnings=[],
            ocr_used=False,
        )
    except Exception as exc:
        if isinstance(exc, DocumentInvalidError):
            raise exc
        raise DocumentInvalidError(f"Failed to parse text document: {str(exc)}")


def parse_document_from_bytes(file_bytes: bytes, file_type: str) -> ExtractedDocument:
    ext = file_type.lower()
    if not ext.startswith("."):
        ext = f".{ext}"
        
    if ext == ".pdf":
        return parse_pdf_bytes(file_bytes)
    elif ext in [".docx", ".doc"]:
        return parse_docx_bytes(file_bytes)
    elif ext in [".txt", ".md", ".csv"]:
        return parse_text_bytes(file_bytes)
    else:
        raise DocumentInvalidError(f"Unsupported file format: {ext}")
